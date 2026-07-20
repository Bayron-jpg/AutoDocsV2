import os
import sys
import threading
import customtkinter
from tkinter import filedialog, messagebox
import docx2pdf
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ------------------ Configuración General ------------------
def obtener_ruta_base():
    if getattr(sys, 'frozen', False):
        return os.path.dirname(sys.executable)  # .exe
    else:
        return os.path.dirname(os.path.abspath(__file__))  # .py
    
BASE_DIR = obtener_ruta_base()
app = customtkinter.CTk()  # Crear ventana vacía
ICON_LIGHT = os.path.join(
    BASE_DIR, "Iconos", "AutoDocsLight.ico"
)  # Icono con fondo blanco
ICON_DARK = os.path.join(
    BASE_DIR, "Iconos", "AutoDocsDark.ico"
)  # Icono con fondo negro
ventana_acerca = None
ventana_plantilla = None
ventana_pdf = None


# ------------------ Funciones ------------------
def crearTexto(app, contenido, letra, tamano):
    label = customtkinter.CTkLabel(
        app,
        text=contenido,  # Texto a ingresar (String)
        font=(
            letra,
            tamano,
        ),  # Tipo de letra (Arial/Roboto/Consolas/etc) y tamaño (12-14-16-etc)
        text_color=("black", "white"),
    )  # Color del texto
    return label


def crearTextoColor(app, contenido, letra, tamano, colorTexto):
    label = customtkinter.CTkLabel(
        app,
        text=contenido,  # Texto a ingresar (String)
        font=(
            letra,
            tamano,
        ),  # Tipo de letra (Arial/Roboto/Consolas/etc) y tamaño (12-14-16-etc)
        text_color=(colorTexto),
    )  # Color del texto
    return label


def cambiar_tema():
    actual = customtkinter.get_appearance_mode()  # Obtener modo actual
    if actual == "Dark":
        customtkinter.set_appearance_mode("Light")
        app.iconbitmap(ICON_LIGHT)
    else:
        customtkinter.set_appearance_mode("Dark")
        app.iconbitmap(ICON_DARK)


# ------------------ Ventanas ------------------
def menu():
    # Inicializar ventana
    customtkinter.set_appearance_mode("Dark")
    app.iconbitmap(ICON_DARK)  # Cambiar icono
    app.title("AutoDocs (1.0.0)")  # Cambiar el tituloDoc superior de la ventana

    # Ajustes de la ventana
    ancho = 800
    alto = 500
    x = (app.winfo_screenwidth() // 2) - (ancho // 2)
    y = (app.winfo_screenheight() // 2) - (alto // 2)
    app.geometry(f"{ancho}x{alto}+{x}+{y}")
    app.resizable(False, False)  # Impide modificar el ancho y el alto

    # ------------ Ajustes personalizados de la ventana ------------
    # --- Posiciones ---
    POS_TITULODOC = (0.24, 0.19)
    POS_TEMA = (0.03, 0.05)
    POS_SALIR = (0.62, 0.76)
    POS_ACERCA = (0.20, 0.76)
    POS_CREAR = (0.12, 0.45)
    POS_CONVERTIR = (0.54, 0.45)

    # Ajustes del Tema (Modo Claro-Oscuro)
    cambiarTema = customtkinter.CTkButton(
        app,
        text="Cambiar Tema",  # Texto del botón
        command=cambiar_tema,  # Función a realizar
        text_color=("black", "white"),  # Color del texto
        fg_color=("#cecece", "gray"),  # Color del botón
        hover_color=("#c0c0c0", "#666666"),
    )  # Color sobre el mouse
    cambiarTema.place(relx=POS_TEMA[0], rely=POS_TEMA[1])

    # TituloDoc AutoDocs
    tituloDoc = crearTexto(app, "Bienvenido(a) a AutoDocs", "Menlo", 36)
    tituloDoc.place(relx=POS_TITULODOC[0], rely=POS_TITULODOC[1])

    # Botón Crear Plantilla
    botonCrear = customtkinter.CTkButton(
        app,
        text="Crear Plantilla",  # Texto del botón
        command=crearPlantilla,  # Función a realizar
        fg_color="#2980b9",  # Color del botón
        hover_color="#1a6a9a",  # Color sobre el mouse
        text_color=("#f0f0f0", "white"),  # Color del texto
        width=270,  # Ancho del botón
        height=110,  # Alto del botón
        font=("Menlo", 24),
    )  # Tipo de letra y tamaño del texto
    botonCrear.place(relx=POS_CREAR[0], rely=POS_CREAR[1])

    # Botón Convertir a PDF
    botonConvertir = customtkinter.CTkButton(
        app,
        text="Convertir a PDF",  # Texto del botón
        command=convertirPdf,  # Función a realizar
        fg_color="#e05a2b",  # Color del botón
        hover_color="#84371B",  # Color sobre el mouse
        text_color=("#f0f0f0", "white"),  # Color del texto
        width=270,  # Ancho del botón
        height=110,  # Alto del botón
        font=("Menlo", 24),
    )  # Tipo de letra y tamaño del texto
    botonConvertir.place(relx=POS_CONVERTIR[0], rely=POS_CONVERTIR[1])

    # Botón Acerca De
    botonAcercaDe = customtkinter.CTkButton(
        app,
        text="Acerca de",  # Texto del botón
        command=acercaDe,  # Función a realizar
        fg_color=("#cecece", "gray"),  # Color del botón
        hover_color=("#c0c0c0", "#666666"),  # Color sobre el mouse
        text_color=("black", "white"),  # Color del texto
        height=35,
    )  # Alto del botón
    botonAcercaDe.place(relx=POS_ACERCA[0], rely=POS_ACERCA[1])

    # Cerrar el programa (Botón de Salir)
    def salir():
        if messagebox.askyesno("Salir", "¿Está seguro(a) que desea salir?"):
            app.destroy()  # Terminar la aplicación

    botonSalir = customtkinter.CTkButton(
        app,
        text="Salir",  # Título de la ventana
        command=salir,  # Función a ejecutar
        fg_color="#ba5858",  # Color del botón
        hover_color="#7f2d2d",  # Color sobre el mouse
        text_color="white",  # Color del texto
        height=35,
    )  # Alto del botón
    botonSalir.place(relx=POS_SALIR[0], rely=POS_SALIR[1])

    app.bind(
        "<Escape>", lambda e: salir()
    )  # Si pulsa Escape pide confirmación para cerrar la aplicación
    app.mainloop()  # Loop que mantiene la ventana abierta


def crearPlantilla():
    global ventana_plantilla

    def verificar_largo(texto_futuro, limite):
        return len(texto_futuro) <= limite

    def estilo(run, size=14, bold=False):
        run.font.name = "Arial"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.bold = bold
    
    from docx.enum.style import WD_STYLE_TYPE

    def estilo_indice(doc):
        estilos_toc = ["TOC 1", "TOC 2", "TOC 3", "Hyperlink"]
        for nombre in estilos_toc:
            try:
                est = doc.styles[nombre]
            except KeyError:
                tipo = WD_STYLE_TYPE.CHARACTER if nombre == "Hyperlink" else WD_STYLE_TYPE.PARAGRAPH
                est = doc.styles.add_style(nombre, tipo)
            est.font.name = "Arial"
            est.font.color.rgb = RGBColor(0, 0, 0)
            est.font.underline = False
            rPr = est.element.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rPr.append(rFonts)
            rFonts.set(qn("w:eastAsia"), "Arial")

    def generarDoc(
        tituloDoc,
        subtituloDoc,
        estudiantes,
        profesor,
        asignatura,
        seccion,
        botonSeleccionarImagen,
        nombreArchivo,
        generarIndice,
        secciones,
    ):
        textoTitulo = tituloDoc.get().strip()
        if not textoTitulo:
            messagebox.showwarning(
                "Campo vacío",
                "Por favor, escriba un título antes de generar el documento.",
            )
            return
        textoSubtitulo = subtituloDoc.get().strip()
        textoEstudiantes = estudiantes.get().strip()
        if not textoEstudiantes:
            messagebox.showwarning(
                "Campo vacío", "Por favor, escriba almenos un estudiante."
            )
            return
        textoProfesor = profesor.get().strip()
        if not textoProfesor:
            messagebox.showwarning(
                "Campo vacío", "Por favor, escriba almenos un profesor(a)."
            )
            return
        textoAsignatura = asignatura.get().strip()
        if not textoAsignatura:
            messagebox.showwarning(
                "Campo vacío", "Por favor, escriba almenos una asignatura."
            )
            return
        textoSeccion = seccion.get().strip()
        nombre = nombreArchivo.get().strip()
        if not nombre:
            nombre = "Documento"
        ruta = os.path.join(BASE_DIR, f"{nombre}.docx")
        doc = Document()

        # === ENCABEZADO (Opcional) ===
        if hasattr(botonSeleccionarImagen, "imagen_path"):
            section = doc.sections[0]
            header = section.header
            header_paragraph = header.paragraphs[0]
            run_header = header_paragraph.add_run()
            run_header.add_picture(
                botonSeleccionarImagen.imagen_path, width=Inches(2.94)
            )
            header_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            header_paragraph.paragraph_format.left_indent = 0
            header_paragraph.paragraph_format.space_before = 0
            header_paragraph.paragraph_format.space_after = 0

        # Eliminar párrafo vacío inicial
        if doc.paragraphs:
            p = doc.paragraphs[0]._element
            p.getparent().remove(p)

        # === Título ===
        titulo = doc.add_heading(textoTitulo, level=0)
        titulo.paragraph_format.space_before = Pt(120)
        titulo.paragraph_format.space_after = Pt(0)
        run = titulo.runs[0]
        run.font.name = "Arial"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        run.font.size = Pt(52)
        run.font.color.rgb = RGBColor(0, 0, 0)
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pPr = titulo._element.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "nil")
        pBdr.append(bottom)
        pPr.append(pBdr)

        # === Subtítulo (Opcional) ===
        if textoSubtitulo:
            parrafo = doc.add_paragraph()
            run_sub = parrafo.add_run(textoSubtitulo)
            estilo(run_sub)
            parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            parrafo.paragraph_format.space_before = Pt(0)
            parrafo.paragraph_format.space_after = Pt(0)

        # === Datos portada ===
        lineas = [
            ("Estudiantes: ", textoEstudiantes),
            ("Profesor(a): ", textoProfesor),
            ("Asignatura: ", textoAsignatura),
        ]
        if textoSeccion:
            lineas.append(("Sección: ", textoSeccion))
        primera = True
        for etiqueta, valor in lineas:
            p = doc.add_paragraph()
            run1 = p.add_run(etiqueta)
            estilo(run1, bold=True)
            run2 = p.add_run(valor)
            estilo(run2)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if primera:
                p.paragraph_format.space_before = Pt(250)
                primera = False
            else:
                p.paragraph_format.space_before = Pt(0)

        # === ÍNDICE (Opcional) ===
        if generarIndice.get():
            estilo_indice(doc)
            titulo_indice = doc.add_heading("Índice", level=1)
            titulo_indice.paragraph_format.page_break_before = True
            titulo_indice.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_indice = titulo_indice.runs[0]
            estilo(run_indice, size=16, bold=True)
            parrafo = doc.add_paragraph()
            run = parrafo.add_run()
            fldChar = OxmlElement("w:fldChar")
            fldChar.set(qn("w:fldCharType"), "begin")
            instrText = OxmlElement("w:instrText")
            instrText.set(qn("xml:space"), "preserve")
            instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
            fldChar2 = OxmlElement("w:fldChar")
            fldChar2.set(qn("w:fldCharType"), "separate")
            fldChar3 = OxmlElement("w:fldChar")
            fldChar3.set(qn("w:fldCharType"), "end")
            run._r.append(fldChar)
            run._r.append(instrText)
            run._r.append(fldChar2)
            run._r.append(fldChar3)

        # === SECCIONES DINÁMICAS ===
        for seccion_data in secciones:
            texto_sec = seccion_data["titulo"].get().strip()
            if not texto_sec:
                continue
            heading = doc.add_heading(texto_sec, level=1)
            heading.paragraph_format.page_break_before = True
            run_h = heading.runs[0]
            estilo(run_h, size=16, bold=True)

            for entry_sub in seccion_data["subtitulos"]:
                texto_sub = entry_sub.get().strip()
                if not texto_sub:
                    continue
                sub = doc.add_heading(texto_sub, level=2)
                run_s = sub.runs[0]
                estilo(run_s, size=14, bold=False)

        doc.save(ruta)
        deseaAbrir = messagebox.askyesno(
            "Documento generado",
            f"El documento se generó correctamente.\n\nRuta:\n{ruta}"
            + (
                "\n\n💡 Para actualizar el índice, abra el documento y presione Ctrl+E / F9"
                if generarIndice.get()
                else ""
            )
            + "\n\n¿Desea abrir el archivo?",
        )
        if deseaAbrir:
            try:
                os.startfile(ruta)
            except Exception as e:
                messagebox.showerror("Error", f"No se pudo abrir el archivo:\n{e}")

    if ventana_plantilla is None or not ventana_plantilla.winfo_exists():
        # ------ Crear ventana ------
        ventana_plantilla = customtkinter.CTkToplevel(app)
        ventana_plantilla.transient(app)
        ventana_plantilla.grab_set()
        ventana_plantilla.focus_force()
        ventana_plantilla.title("Creando Plantilla")
        ventana_plantilla.resizable(False, False)

        ancho = 800
        alto = 500
        x = app.winfo_x() + (app.winfo_width() - ancho) // 2
        y = app.winfo_y() + (app.winfo_height() - alto) // 2
        ventana_plantilla.geometry(f"{ancho}x{alto}+{x}+{y}")
        ventana_plantilla.update_idletasks()
        ventana_plantilla.bind("<Escape>", lambda e: ventana_plantilla.destroy())

        icono = (
            ICON_DARK if customtkinter.get_appearance_mode() == "Dark" else ICON_LIGHT
        )
        ventana_plantilla.after(200, lambda: ventana_plantilla.iconbitmap(icono))

        # --- Validaciones ---
        cmd_titulo = (
            ventana_plantilla.register(lambda p: verificar_largo(p, 60)),
            "%P",
        )
        cmd_subtitulo = (
            ventana_plantilla.register(lambda p: verificar_largo(p, 125)),
            "%P",
        )
        cmd_estudiantes = (
            ventana_plantilla.register(lambda p: verificar_largo(p, 50)),
            "%P",
        )
        cmd_profesor = (
            ventana_plantilla.register(lambda p: verificar_largo(p, 44)),
            "%P",
        )
        cmd_asignatura = (
            ventana_plantilla.register(lambda p: verificar_largo(p, 44)),
            "%P",
        )
        cmd_seccion = (
            ventana_plantilla.register(lambda p: verificar_largo(p, 30)),
            "%P",
        )

        # ==================== FRAME 1 (Portada) ====================
        frame1 = customtkinter.CTkFrame(ventana_plantilla, fg_color="transparent")
        frame1.place(relx=0, rely=0, relwidth=1, relheight=1)

        def ir_a_pagina2():
            frame1.place_forget()
            frame2.place(relx=0, rely=0, relwidth=1, relheight=1)

        def ir_a_pagina1():
            frame2.place_forget()
            frame1.place(relx=0, rely=0, relwidth=1, relheight=1)

        crearTexto(frame1, "Crear Plantilla de Documento", "Menlo", 28).place(
            relx=0.27, rely=0.06
        )

        crearTexto(frame1, "1.Título del documento", "Consolas", 16).place(
            relx=0.10, rely=0.20
        )
        tituloDoc = customtkinter.CTkEntry(
            frame1,
            placeholder_text="Escribir título...",
            width=200,
            height=35,
            validate="key",
            validatecommand=cmd_titulo,
        )
        tituloDoc.place(relx=0.10, rely=0.26)

        crearTexto(frame1, "2.Subtítulo", "Consolas", 16).place(relx=0.38, rely=0.20)
        subtituloDoc = customtkinter.CTkEntry(
            frame1,
            placeholder_text="(Opcional)",
            width=200,
            height=35,
            validate="key",
            validatecommand=cmd_subtitulo,
        )
        subtituloDoc.place(relx=0.38, rely=0.26)

        crearTexto(frame1, "3.Estudiantes", "Consolas", 16).place(relx=0.66, rely=0.20)
        estudiantes = customtkinter.CTkEntry(
            frame1,
            placeholder_text="Ej: Pedro - Juan - Sofía",
            width=200,
            height=35,
            validate="key",
            validatecommand=cmd_estudiantes,
        )
        estudiantes.place(relx=0.66, rely=0.26)

        crearTexto(frame1, "4.Profesor(a)", "Consolas", 16).place(relx=0.10, rely=0.35)
        profesor = customtkinter.CTkEntry(
            frame1,
            placeholder_text="Nombre del profesor(a)",
            width=200,
            height=35,
            validate="key",
            validatecommand=cmd_profesor,
        )
        profesor.place(relx=0.10, rely=0.41)

        crearTexto(frame1, "5.Asignatura", "Consolas", 16).place(relx=0.38, rely=0.35)
        asignatura = customtkinter.CTkEntry(
            frame1,
            placeholder_text="Ej: Fundamentos de Software",
            width=200,
            height=35,
            validate="key",
            validatecommand=cmd_asignatura,
        )
        asignatura.place(relx=0.38, rely=0.41)

        crearTexto(frame1, "6.Sección", "Consolas", 16).place(relx=0.66, rely=0.35)
        seccion = customtkinter.CTkEntry(
            frame1,
            placeholder_text="(Opcional) Ej: 008D",
            width=200,
            height=35,
            validate="key",
            validatecommand=cmd_seccion,
        )
        seccion.place(relx=0.66, rely=0.41)

        crearTexto(frame1, "Seleccionar imagen de encabezado", "Consolas", 16).place(
            relx=0.14, rely=0.55
        )
        labelImagen = crearTexto(frame1, "(Opcional)", "Consolas", 12)
        labelImagen.place(relx=0.33, rely=0.63)

        def seleccionarImagen():
            imagen = filedialog.askopenfilename(
                title="Seleccionar imagen de encabezado",
                filetypes=[("Imágenes", "*.png *.jpg *.jpeg")],
            )
            if imagen:
                labelImagen.configure(text=os.path.basename(imagen))
                botonSeleccionarImagen.imagen_path = imagen

        botonSeleccionarImagen = customtkinter.CTkButton(
            frame1,
            text="Explorar Archivos",
            command=seleccionarImagen,
            fg_color=("#cecece", "gray"),
            hover_color=("#c0c0c0", "#666666"),
            text_color=("black", "white"),
            height=35,
        )
        botonSeleccionarImagen.place(relx=0.14, rely=0.63)

        customtkinter.CTkButton(
            frame1,
            text="Volver",
            command=ventana_plantilla.destroy,
            fg_color="#ba6258",
            hover_color="#7f342d",
            text_color="white",
            height=35,
        ).place(relx=0.10, rely=0.83)

        customtkinter.CTkButton(
            frame1,
            text="Datos de Prueba",
            command=lambda: rellenarPrueba(),
            fg_color=("#cecece", "gray"),
            hover_color=("#c0c0c0", "#666666"),
            text_color=("black", "white"),
            height=35,
        ).place(relx=0.38, rely=0.83)

        customtkinter.CTkButton(
            frame1,
            text="Siguiente",
            command=ir_a_pagina2,
            fg_color="#437791",
            hover_color="#386379",
            text_color="white",
            height=35,
            width=150,
        ).place(relx=0.71, rely=0.83)

        # ==================== FRAME 2 (Contenido + Exportar) ====================
        frame2 = customtkinter.CTkFrame(ventana_plantilla, fg_color="transparent")

        crearTexto(frame2, "Contenido del Documento", "Menlo", 28).place(
            relx=0.30, rely=0.05
        )

        # --- Checkbox índice ---
        generarIndice = customtkinter.CTkCheckBox(
            frame2,
            text="Generar índice automático en página 2",
            font=("Consolas", 13),
            text_color=("black", "white"),
        )
        generarIndice.place(relx=0.31, rely=0.17)

        # --- Nombre del archivo ---
        crearTexto(frame2, "Nombre del archivo", "Consolas", 14).place(
            relx=0.39, rely=0.81
        )
        nombreArchivo = customtkinter.CTkEntry(
            frame2, placeholder_text="(Opcional) Ej: Informe N°1", width=230, height=35
        )
        nombreArchivo.place(relx=0.34, rely=0.87)

        # --- Secciones dinámicas ---
        crearTexto(frame2, "Secciones del documento", "Consolas", 14).place(
            relx=0.10, rely=0.28
        )

        secciones = []

        scroll_secciones = customtkinter.CTkScrollableFrame(
            frame2, width=730, height=185
        )
        scroll_secciones.place(relx=0.03, rely=0.35)

        def agregar_seccion():
            idx = len(secciones) + 1
            frame_sec = customtkinter.CTkFrame(
                scroll_secciones, fg_color=("#d9d9d9", "#3a3a3a"), corner_radius=3
            )
            frame_sec.pack(fill="x", padx=2, pady=1)

            scroll_secciones.update_idletasks()

            customtkinter.CTkLabel(
                frame_sec,
                text=f"Sección {idx}:",
                font=("Consolas", 13),
                text_color=("black", "white"),
            ).grid(row=0, column=0, padx=8, pady=2, sticky="w")

            entry_titulo = customtkinter.CTkEntry(
                frame_sec, placeholder_text="Título de la sección", width=280, height=26
            )
            entry_titulo.grid(row=0, column=1, padx=5, pady=2)

            subtitulos = []

            frame_subs = None

            def agregar_subtitulo():
                nonlocal frame_subs

                if frame_subs is None:
                    frame_subs = customtkinter.CTkFrame(
                        frame_sec, fg_color="transparent"
                    )
                    frame_subs.grid(
                        row=1, column=0, columnspan=4, padx=15, pady=2, sticky="w"
                    )

                sub_idx = len(subtitulos) + 1

                frame_sub = customtkinter.CTkFrame(frame_subs, fg_color="transparent")
                frame_sub.pack(fill="x", pady=1)

                customtkinter.CTkLabel(
                    frame_sub,
                    text=f"  └ Subtítulo {sub_idx}:",
                    font=("Consolas", 12),
                    text_color=("black", "white"),
                ).pack(side="left", padx=4)

                entry_sub = customtkinter.CTkEntry(
                    frame_sub, placeholder_text="Subtítulo...", width=220, height=28
                )
                entry_sub.pack(side="left", padx=4)

                def eliminar_subtitulo():
                    subtitulos.remove(entry_sub)
                    frame_sub.destroy()

                customtkinter.CTkButton(
                    frame_sub,
                    text="✕",
                    width=28,
                    height=28,
                    fg_color="#ba6258",
                    hover_color="#7f342d",
                    text_color="white",
                    command=eliminar_subtitulo,
                ).pack(side="left", padx=2)

                subtitulos.append(entry_sub)

            def eliminar_seccion():
                secciones.remove(seccion_data)
                frame_sec.destroy()

            customtkinter.CTkButton(
                frame_sec,
                text="+ Subtítulo",
                width=110,
                height=28,
                fg_color=("#cecece", "gray"),
                hover_color=("#c0c0c0", "#666666"),
                text_color=("black", "white"),
                command=agregar_subtitulo,
            ).grid(row=0, column=2, padx=8)

            customtkinter.CTkButton(
                frame_sec,
                text="✕ Sección",
                width=100,
                height=28,
                fg_color="#ba6258",
                hover_color="#7f342d",
                text_color="white",
                command=eliminar_seccion,
            ).grid(row=0, column=3, padx=4)

            seccion_data = {
                "titulo": entry_titulo,
                "subtitulos": subtitulos,
                "frame": frame_sec,
            }
            secciones.append(seccion_data)

        customtkinter.CTkButton(
            frame2,
            text="+ Agregar Sección",
            fg_color=("#cecece", "gray"),
            hover_color=("#c0c0c0", "#666666"),
            text_color=("black", "white"),
            height=30,
            width=160,
            command=agregar_seccion,
        ).place(relx=0.76, rely=0.28)

        customtkinter.CTkButton(
            frame2,
            text="Anterior",
            command=ir_a_pagina1,
            fg_color=("#cecece", "gray"),
            hover_color=("#c0c0c0", "#666666"),
            text_color=("black", "white"),
            height=35,
        ).place(relx=0.10, rely=0.88)

        customtkinter.CTkButton(
            frame2,
            text="Generar Plantilla",
            command=lambda: generarDoc(
                tituloDoc,
                subtituloDoc,
                estudiantes,
                profesor,
                asignatura,
                seccion,
                botonSeleccionarImagen,
                nombreArchivo,
                generarIndice,
                secciones,
            ),
            fg_color="#437791",
            hover_color="#386379",
            text_color="white",
            height=40,
            width=170,
        ).place(relx=0.69, rely=0.86)

        # --- Función datos de prueba ---
        def rellenarPrueba():
            tituloDoc.delete(0, "end")
            tituloDoc.insert(0, "Informe Gestión de Biblioteca")
            subtituloDoc.delete(0, "end")
            subtituloDoc.insert(0, "Etapa 1: Análisis de Requisitos")
            estudiantes.delete(0, "end")
            estudiantes.insert(0, "Bayron Urrutia - José Ibarra")
            profesor.delete(0, "end")
            profesor.insert(0, "Elliana Mallén González")
            asignatura.delete(0, "end")
            asignatura.insert(0, "Ingeniería de Requisitos")
            seccion.delete(0, "end")
            seccion.insert(0, "008D")
            nombreArchivo.delete(0, "end")
            nombreArchivo.insert(0, "Informe_Biblioteca")

    else:
        ventana_plantilla.focus()


def convertirPdf():
    # --- Posiciones ---
    POS_TITULODOC = (0.18, 0.09)
    POS_ARCHIVO = (0.40, 0.44)
    POS_SELECCIONAR = (0.19, 0.43)
    POS_VOLVER = (0.41, 0.83)
    POS_CONVERTIR = (0.41, 0.56)
    POS_TEXTOARCHIVO = (0.19, 0.34)

    global ventana_pdf

    if ventana_pdf is None or not ventana_pdf.winfo_exists():
        # ------ Crear y ajustar ventana ------
        ventana_pdf = customtkinter.CTkToplevel(app)
        ventana_pdf.transient(app)  # Vincula a la ventana principal
        ventana_pdf.grab_set()  # Impide interactuar con la principal
        ventana_pdf.focus_force()  # Le da foco inmediatamente
        ventana_pdf.title(
            "Convertir Documento Word a PDF"
        )  # Cambiar el tituloDoc superior de la ventana
        ventana_pdf.resizable(False, False)  # Impide agrandar o achicar la ventana

        # Ajustes de la ventana
        ancho = 800
        alto = 500
        x = app.winfo_x() + (app.winfo_width() - ancho) // 2
        y = app.winfo_y() + (app.winfo_height() - alto) // 2
        ventana_pdf.geometry(f"{ancho}x{alto}+{x}+{y}")

        # Fuerza a que la ventana termine de crearse para aplicar el icono
        ventana_pdf.update_idletasks()
        ventana_pdf.bind(
            "<Escape>", lambda e: ventana_pdf.destroy()
        )  # Si pulsa Escape se cierra la ventana

        # Obtener el modo actual (Claro/Oscuro)
        icono = (
            ICON_DARK if customtkinter.get_appearance_mode() == "Dark" else ICON_LIGHT
        )
        ventana_pdf.after(
            200, lambda: ventana_pdf.iconbitmap(icono)
        )  # Espera 200ms para cambiar el icono

        # --- Título ---
        tituloDoc = crearTexto(
            ventana_pdf, "Convertir Documento Word a PDF", "Menlo", 34
        )
        tituloDoc.place(relx=POS_TITULODOC[0], rely=POS_TITULODOC[1])

        # -- Texto seleccion de archivo ---
        textoArchivo = crearTexto(
            ventana_pdf, "Seleccionar archivo PDF a convertir", "Consolas", 16
        )
        textoArchivo.place(relx=POS_TEXTOARCHIVO[0], rely=POS_TEXTOARCHIVO[1])

        # --- Archivo seleccionado ---
        archivoSeleccionado = crearTexto(
            ventana_pdf, "Ningún archivo seleccionado", "Consolas", 12
        )
        archivoSeleccionado.place(relx=POS_ARCHIVO[0], rely=POS_ARCHIVO[1])

        # --- Función seleccionar archivo ---
        def seleccionarArchivo():
            archivo = filedialog.askopenfilename(
                title="Seleccionar documento Word",
                filetypes=[("Documentos Word", "*.docx *.doc")],
            )
            if archivo:
                archivoSeleccionado.configure(text=os.path.basename(archivo))
                botonConvertir.configure(state="normal")
                botonConvertir.archivo_path = archivo

        # --- Función convertir ---
        def convertir():
            archivo = botonConvertir.archivo_path

            # Deshabilitar botones mientras convierte
            botonConvertir.configure(state="disabled", text="Convirtiendo...")
            botonSeleccionar.configure(state="disabled")

            def proceso():
                try:
                    docx2pdf.convert(archivo)
                    ruta_pdf = os.path.splitext(archivo)[0] + ".pdf"
                    ventana_pdf.after(0, lambda: exito(ruta_pdf))
                except Exception as e:
                    ventana_pdf.after(0, lambda: error(str(e)))

            def exito(ruta):
                botonConvertir.configure(
                    state="disabled", text="Convertir a PDF"
                )  # Deshabilitar y resetear texto
                botonSeleccionar.configure(state="normal")
                archivoSeleccionado.configure(
                    text="Ningún archivo seleccionado"
                )  # Vaciar label
                messagebox.showinfo(
                    "PDF Generado satisfactoriamente",
                    f"PDF generado correctamente.\n\nGuardado en:\n{ruta}",
                )

            def error(e):
                botonConvertir.configure(state="normal", text="Convertir a PDF")
                botonSeleccionar.configure(state="normal")
                messagebox.showerror("Error", f"No se pudo convertir:\n{e}")

            threading.Thread(target=proceso, daemon=True).start()

        # --- Botón seleccionar ---
        botonSeleccionar = customtkinter.CTkButton(
            ventana_pdf,
            text="Seleccionar archivo",
            command=seleccionarArchivo,
            fg_color=("#cecece", "gray"),
            hover_color=("#c0c0c0", "#666666"),
            text_color=("black", "white"),
            height=35,
        )
        botonSeleccionar.place(relx=POS_SELECCIONAR[0], rely=POS_SELECCIONAR[1])

        # --- Botón convertir (deshabilitado hasta seleccionar archivo) ---
        botonConvertir = customtkinter.CTkButton(
            ventana_pdf,
            text="Convertir a PDF",
            command=convertir,
            fg_color="#437791",
            hover_color="#386379",
            text_color="white",
            height=35,
            state="disabled",
        )
        botonConvertir.place(relx=POS_CONVERTIR[0], rely=POS_CONVERTIR[1])

        botonVolver = customtkinter.CTkButton(
            ventana_pdf,
            text="Volver",  # Título de la ventana
            command=ventana_pdf.destroy,  # Función a ejecutar
            fg_color="#ba6258",  # Color del botón
            hover_color="#7f342d",  # Color sobre el mouse
            text_color="white",  # Color del texto
            height=35,
        )  # Alto del botón
        botonVolver.place(relx=POS_VOLVER[0], rely=POS_VOLVER[1])

    else:
        ventana_pdf.focus()


def acercaDe():
    # --- Posiciones ---
    POS_AUTODOCS = (0.35, 0.03)
    POS_GENERARPLANTILLA = (0.07, 0.21)
    POS_SALIR = (0.32, 0.83)

    global ventana_acerca

    if ventana_acerca is None or not ventana_acerca.winfo_exists():
        # ------ Crear y ajustar ventana ------
        ventana_acerca = customtkinter.CTkToplevel(app)
        ventana_acerca.transient(app)  # Vincula a la ventana principal
        ventana_acerca.grab_set()  # Impide interactuar con la principal
        ventana_acerca.focus_force()  # Le da foco inmediatamente
        ventana_acerca.title("Acerca de")  # Cambiar el tituloDoc superior de la ventana
        ventana_acerca.resizable(False, False)  # Impide agrandar o achicar la ventana

        # Ajustes de la ventana
        ancho = 400
        alto = 500
        x = app.winfo_x() + (app.winfo_width() - ancho) // 2
        y = app.winfo_y() + (app.winfo_height() - alto) // 2
        ventana_acerca.geometry(f"{ancho}x{alto}+{x}+{y}")

        # Fuerza a que la ventana termine de crearse para aplicar el icono
        ventana_acerca.update_idletasks()
        ventana_acerca.bind(
            "<Escape>", lambda e: ventana_acerca.destroy()
        )  # Si pulsa Escape se cierra la ventana

        # Obtener el modo actual (Claro/Oscuro)
        icono = (
            ICON_DARK if customtkinter.get_appearance_mode() == "Dark" else ICON_LIGHT
        )
        ventana_acerca.after(
            200, lambda: ventana_acerca.iconbitmap(icono)
        )  # Espera 200ms para cambiar el icono

        autoDocs = crearTexto(ventana_acerca, "AutoDocs", "Consolas", 28)
        autoDocs.place(relx=POS_AUTODOCS[0], rely=POS_AUTODOCS[1])

        texto = customtkinter.CTkLabel(
            ventana_acerca,
            text="Aplicación para automatizar la creación\n"
            "de documentos Word repetitivos.\n\n"
            "Proyecto personal desarrollado en\n"
            "Python con CustomTkinter, Python-Docx y\n"
            "Docx2Pdf.\n\n\n"
            "Desarrollado por:\n"
            "Bayron Urrutia\n\n"
            "© 2026",
            font=("Consolas", 16),  # Tamaño y tipo de letra
            justify="left",
        )  # Alineamiento del texto a la izquierda
        texto.place(relx=POS_GENERARPLANTILLA[0], rely=POS_GENERARPLANTILLA[1])

        botonSalir = customtkinter.CTkButton(
            ventana_acerca,
            text="Volver",  # Título de la ventana
            command=ventana_acerca.destroy,  # Función a ejecutar
            fg_color="#ba6258",  # Color del botón
            hover_color="#7f342d",  # Color sobre el mouse
            text_color="white",  # Color del texto
            height=35,
        )  # Alto del botón
        botonSalir.place(relx=POS_SALIR[0], rely=POS_SALIR[1])

    else:
        ventana_acerca.focus()


# ------------------ Inicio ------------------
menu()
